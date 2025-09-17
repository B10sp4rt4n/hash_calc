from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def generar_acta_docx_bytes(meta: dict, hashes: dict, encadenamiento: dict | None = None):
    """
    Devuelve (filename, bytes) de un Acta DOCX en memoria.
    """
    doc = Document()

    def add_title(text, size=18, bold=True):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_title("ACTA DE CONSTANCIA DE INTEGRIDAD")
    add_title("Cadena de confianza documental — Semilla AUP", size=12, bold=False)
    doc.add_paragraph("")

    meta_lines = [
        f"Documento referenciado (título): {meta.get('titulo_documento','')}",
        f"Autor custodio: {meta.get('actor','')}",
        f"Lugar: {meta.get('lugar','')}",
        f"Fecha del acta: {meta.get('fecha_acta','')}",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph("")
    doc.add_paragraph("1) Identificación del artefacto cerrado:")
    doc.add_paragraph(f"   • Nombre de archivo: {meta.get('nombre_archivo','')}")
    doc.add_paragraph(f"   • Tamaño (bytes): {meta.get('tamano_bytes','')}")
    doc.add_paragraph(f"   • Tipo/MIME: {meta.get('mimetype','')}")
    if meta.get("notas"):
        doc.add_paragraph(f"   • Observaciones: {meta.get('notas','')}")

    doc.add_paragraph("")
    doc.add_paragraph("2) Huella(s) criptográfica(s):")
    for k,v in hashes.items():
        if v:
            doc.add_paragraph(f"   • {k.upper()}: {v}")

    doc.add_paragraph("")
    doc.add_paragraph("3) Encadenamiento (opcional):")
    if encadenamiento and encadenamiento.get("hash_previo"):
        doc.add_paragraph(f"   • Hash del acta/artefacto previo: {encadenamiento['hash_previo']}")
    else:
        doc.add_paragraph("   • (sin datos)")

    doc.add_paragraph("")
    doc.add_paragraph("4) Firmas:")
    doc.add_paragraph("   • Autor/Custodio: ____________________    Firma: ____________________    Fecha: ____________________")
    doc.add_paragraph("   • Testigo (opcional): ____________________    Firma: ____________________    Fecha: ____________________")

    doc.add_paragraph("")
    doc.add_paragraph("5) Bitácora forense (referencia):")
    doc.add_paragraph("   • ID de registro: ____________________")
    doc.add_paragraph("   • Enlace o folio interno: ____________________")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"Acta_Integridad_{meta.get('nombre_archivo','documento')}.docx".replace(" ", "_")
    return fname, buf.getvalue()